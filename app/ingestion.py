"""
app/ingestion.py — пайплайн загрузки и индексации документов.

Этапы:
1. Парсинг PDF (PyMuPDF + section title extraction) / DOCX (python-docx) / TXT.
   Опционально Docling — если установлен и BAAI/bge-m3 не используется для парсинга.
2. Parent-child chunking: большие parent-чанки для генерации, маленькие child — для поиска.
3. Embedding через bge-m3 (TEI HTTP или локально через sentence-transformers).
4. Запись в Qdrant с named vectors (dense + sparse reserved) и rich payload.
5. Snapshot после индекски (для backup).
6. Идемпотентное удаление по doc_id (filter).
7. Дедупликация по chunk_hash (опционально — пока только сохраняется).
"""

import logging
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from llama_index.core.node_parser import SentenceSplitter
from qdrant_client import QdrantClient
from qdrant_client.http.models import (
    Distance,
    FieldCondition,
    Filter,
    MatchValue,
    PayloadSchemaType,
    PointStruct,
    SparseIndexParams,
    SparseVectorParams,
    VectorParams,
)

from app.config import settings
from app.utils import (
    child_id_for,
    compute_chunk_hash,
    detect_language,
    extract_section_title_from_pdf_page,
    generate_doc_id,
    parent_id_for,
    sanitize_text,
    utc_now_iso,
)

logger = logging.getLogger(__name__)

# Опциональный импорт Docling
try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    logger.info("Docling не установлен — используется PyMuPDF/python-docx как fallback.")


# ---------------------------------------------------------------------------
# Qdrant client (singleton)
# ---------------------------------------------------------------------------
_qdrant_client: Optional[QdrantClient] = None


def get_qdrant_client() -> QdrantClient:
    global _qdrant_client
    if _qdrant_client is None:
        if settings.QDRANT_PATH:
            # embedded-режим: локальное файловое хранилище, без сервера/Docker
            _qdrant_client = QdrantClient(path=settings.QDRANT_PATH)
        else:
            _qdrant_client = QdrantClient(
                host=settings.QDRANT_HOST,
                port=settings.QDRANT_PORT,
                timeout=120.0,
            )
    return _qdrant_client


def ensure_collection() -> None:
    """Создаёт коллекцию с named vectors (dense + reserved sparse) + payload indexes."""
    client = get_qdrant_client()
    existing = {c.name for c in client.get_collections().collections}

    if settings.QDRANT_COLLECTION not in existing:
        client.create_collection(
            collection_name=settings.QDRANT_COLLECTION,
            vectors_config={
                "dense": VectorParams(
                    size=settings.EMBED_DIM,
                    distance=Distance.COSINE,
                ),
            },
            sparse_vectors_config={
                "sparse": SparseVectorParams(
                    index=SparseIndexParams(on_disk=False)
                ),
            },
        )
        logger.info("Created collection '%s' (dense=%d, sparse reserved).",
                    settings.QDRANT_COLLECTION, settings.EMBED_DIM)
    _ensure_payload_indexes(client)


def _ensure_payload_indexes(client: QdrantClient) -> None:
    """Создаёт payload-индексы для быстрой фильтрации и delete-by-filter."""
    indexes = [
        ("doc_id", PayloadSchemaType.KEYWORD),
        ("language", PayloadSchemaType.KEYWORD),
        ("source_type", PayloadSchemaType.KEYWORD),
        ("page_number", PayloadSchemaType.INTEGER),
        ("doc_version", PayloadSchemaType.KEYWORD),
        ("parent_id", PayloadSchemaType.KEYWORD),
        ("chunk_hash", PayloadSchemaType.KEYWORD),
    ]
    for field, schema in indexes:
        try:
            client.create_payload_index(
                collection_name=settings.QDRANT_COLLECTION,
                field_name=field,
                field_schema=schema,
            )
        except Exception:
            # Уже существует — игнорируем
            pass
    logger.info("Payload indexes ensured.")


def create_snapshot() -> Optional[str]:
    """Создаёт snapshot коллекции для backup."""
    try:
        client = get_qdrant_client()
        snapshot = client.create_snapshot(collection_name=settings.QDRANT_COLLECTION)
        logger.info("Snapshot created: %s", snapshot.name)
        return snapshot.name
    except Exception as exc:
        logger.warning("Snapshot failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# Парсеры
# ---------------------------------------------------------------------------

def parse_pdf(path: Path) -> List[Dict[str, Any]]:
    """
    PyMuPDF: текст + номер страницы + заголовок раздела (bold-эвристика).
    """
    pages: List[Dict[str, Any]] = []
    with fitz.open(str(path)) as doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""
            text = sanitize_text(text)
            if not text or len(text) < 30:
                continue
            section = extract_section_title_from_pdf_page(page)
            pages.append({
                "text": text,
                "page_number": page_num,
                "section_title": section,
                "source_type": "pdf",
            })
    logger.info("PDF '%s': parsed %d pages.", path.name, len(pages))
    return pages


def parse_docx(path: Path) -> List[Dict[str, Any]]:
    """
    python-docx: параграфы + таблицы. Заголовок раздела берётся из последнего Heading.
    """
    doc = DocxDocument(str(path))
    current_section = ""
    paragraphs: List[str] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            current_section = text[:120]
            paragraphs.append(f"\n\n## {text}\n")
        else:
            paragraphs.append(text)

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    full_text = sanitize_text("\n".join(paragraphs))
    if not full_text:
        return []
    return [{
        "text": full_text,
        "page_number": 1,   # DOCX не имеет страниц
        "section_title": current_section,
        "source_type": "docx",
    }]


def parse_txt(path: Path) -> List[Dict[str, Any]]:
    """TXT с авто-определением кодировки."""
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            text = path.read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")

    text = sanitize_text(text)
    if not text:
        return []
    return [{
        "text": text,
        "page_number": 1,
        "section_title": "",
        "source_type": "txt",
    }]


def parse_document(path: Path) -> List[Dict[str, Any]]:
    """Маршрутизация по расширению."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".txt":
        return parse_txt(path)
    raise ValueError(f"Неподдерживаемый формат: {ext}")


# ---------------------------------------------------------------------------
# Parent-child chunking
# ---------------------------------------------------------------------------

def get_parent_splitter() -> SentenceSplitter:
    return SentenceSplitter(
        chunk_size=settings.PARENT_CHUNK_SIZE,
        chunk_overlap=settings.PARENT_CHUNK_OVERLAP,
        paragraph_separator="\n\n",
        secondary_chunking_regex="[.!?]\\s+",
    )


def get_child_splitter() -> SentenceSplitter:
    return SentenceSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        paragraph_separator="\n\n",
        secondary_chunking_regex="[.!?]\\s+",
    )


def chunk_document(
    pages: List[Dict[str, Any]],
    doc_id: str,
    doc_name: str,
    source_path: str,
) -> List[Dict[str, Any]]:
    """
    Parent-child chunking: каждый page → parent chunks → child chunks.
    Сохраняет rich metadata в каждом child.
    """
    parent_splitter = get_parent_splitter()
    child_splitter = get_child_splitter()

    nodes: List[Dict[str, Any]] = []
    paragraph_counter = 0

    for page in pages:
        page_text = page["text"]
        if not page_text:
            continue

        page_num = page["page_number"]
        section_title = page.get("section_title", "")
        source_type = page.get("source_type", "")

        try:
            parent_chunks = parent_splitter.split_text(page_text)
        except Exception as exc:
            logger.warning("Parent split failed for page %s: %s", page_num, exc)
            parent_chunks = [page_text]

        for p_idx, parent_text in enumerate(parent_chunks):
            parent_text = parent_text.strip()
            if not parent_text:
                continue

            parent_pid = parent_id_for(doc_id, p_idx)
            parent_chunk_hash = compute_chunk_hash(parent_text)

            # Parent-точка (для PDR — большие чанки)
            nodes.append({
                "text": parent_text,
                "doc_id": doc_id,
                "doc_name": doc_name,
                "source_path": source_path,
                "doc_version": "v1.0",
                "page_number": page_num,
                "paragraph_index": paragraph_counter,
                "section_title": section_title or f"Страница {page_num}",
                "section_path": str(page_num),
                "parent_id": None,         # сам parent
                "is_parent": True,
                "chunk_id": parent_pid,
                "chunk_hash": parent_chunk_hash,
                "language": detect_language(parent_text),
                "source_type": source_type,
                "created_at": utc_now_iso(),
            })
            paragraph_counter += 1

            # Child chunks
            try:
                child_chunks = child_splitter.split_text(parent_text)
            except Exception as exc:
                logger.warning("Child split failed: %s", exc)
                child_chunks = [parent_text]

            for c_idx, child_text in enumerate(child_chunks):
                child_text = child_text.strip()
                if not child_text:
                    continue

                child_cid = child_id_for(doc_id, p_idx, c_idx)
                nodes.append({
                    "text": child_text,
                    "doc_id": doc_id,
                    "doc_name": doc_name,
                    "source_path": source_path,
                    "doc_version": "v1.0",
                    "page_number": page_num,
                    "paragraph_index": paragraph_counter,
                    "section_title": section_title or f"Страница {page_num}",
                    "section_path": str(page_num),
                    "parent_id": parent_pid,
                    "is_parent": False,
                    "chunk_id": child_cid,
                    "chunk_hash": compute_chunk_hash(child_text),
                    "language": detect_language(child_text),
                    "source_type": source_type,
                    "created_at": utc_now_iso(),
                })
                paragraph_counter += 1

    logger.info("Doc '%s': %d nodes (parent + child).", doc_name, len(nodes))
    return nodes


# ---------------------------------------------------------------------------
# Embedding
# ---------------------------------------------------------------------------

_embed_model = None


def get_embedder():
    """Singleton sentence-transformers для bge-m3."""
    global _embed_model
    if _embed_model is None:
        from sentence_transformers import SentenceTransformer
        logger.info("Loading embedding model: %s (device=%s)", settings.EMBED_MODEL, settings.EMBED_DEVICE)
        _embed_model = SentenceTransformer(
            settings.EMBED_MODEL,
            device=settings.EMBED_DEVICE,
            cache_folder=settings.MODELS_DIR,
        )
        logger.info("Embedding model loaded (dim=%d).", settings.EMBED_DIM)
    return _embed_model


def embed_texts(texts: List[str]) -> List[List[float]]:
    """
    Возвращает dense embeddings (1024-dim для bge-m3).
    Если USE_TEI=True — ходит по HTTP в tei-embeddings сервис.
    """
    if not texts:
        return []

    if settings.USE_TEI:
        import httpx
        with httpx.Client(timeout=60.0) as client:
            resp = client.post(
                f"{settings.tei_url}/embed",
                json={"inputs": texts},
            )
            resp.raise_for_status()
            return resp.json()

    embedder = get_embedder()
    embeddings = embedder.encode(texts, normalize_embeddings=True, show_progress_bar=False)
    return embeddings.tolist()


# ---------------------------------------------------------------------------
# Qdrant operations
# ---------------------------------------------------------------------------

def _delete_doc_points(client: QdrantClient, doc_id: str) -> None:
    """Удаляет все точки (parent + child) конкретного документа."""
    client.delete(
        collection_name=settings.QDRANT_COLLECTION,
        points_selector=Filter(
            must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
        ),
    )


def ingest_document(path: Path) -> Dict[str, Any]:
    """
    Полный пайплайн: parse → parent-child chunk → embed → upsert → snapshot.
    Идемпотентно: повторная загрузка того же файла удаляет старые точки.
    """
    ensure_collection()

    doc_name = path.name
    doc_id = generate_doc_id(path)
    logger.info("Ingesting '%s' (doc_id=%s)", doc_name, doc_id)

    pages = parse_document(path)
    if not pages:
        raise ValueError(f"Документ пуст или не распознан: {doc_name}")

    nodes = chunk_document(pages, doc_id, doc_name, str(path))
    if not nodes:
        raise ValueError(f"Не удалось получить чанки из: {doc_name}")

    client = get_qdrant_client()

    # Идемпотентность: удаляем старые точки этого документа
    try:
        _delete_doc_points(client, doc_id)
        logger.info("Pre-deleted old points for doc_id=%s", doc_id)
    except Exception as exc:
        logger.warning("Pre-delete failed: %s", exc)

    # Дедупликация по chunk_hash (взят из v2-fixed)
    # Если тот же chunk_hash уже есть в коллекции (например, при повторной загрузке
    # того же файла) — пропускаем, чтобы не плодить дубликаты.
    skipped_duplicates = 0
    if settings.USE_CHUNK_HASH_DEDUP:
        existing_hashes = _get_existing_chunk_hashes(client, doc_id)
        original_count = len(nodes)
        deduped: List[Dict[str, Any]] = []
        seen_in_this_run: set = set()
        for n in nodes:
            h = n.get("chunk_hash")
            if h and (h in existing_hashes or h in seen_in_this_run):
                skipped_duplicates += 1
                continue
            seen_in_this_run.add(h)
            deduped.append(n)
        nodes = deduped
        if skipped_duplicates > 0:
            logger.info("Skipped %d duplicate chunks (by chunk_hash).", skipped_duplicates)

    if not nodes:
        logger.info("All chunks were duplicates — nothing to ingest.")
        return {
            "doc_id": doc_id,
            "doc_name": doc_name,
            "pages": len(pages),
            "parents": 0,
            "chunks": 0,
            "total_nodes": 0,
            "snapshot": None,
            "skipped_duplicates": skipped_duplicates,
        }

    # Embedding всех текстов одним батчем (для скорости)
    texts = [n["text"] for n in nodes]
    batch_size = 32
    all_embeddings: List[List[float]] = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        embs = embed_texts(batch)
        all_embeddings.extend(embs)
        logger.info("Embedded batch %d/%d", i // batch_size + 1, (len(texts) + batch_size - 1) // batch_size)

    # Формируем points для Qdrant
    points: List[PointStruct] = []
    for i, node in enumerate(nodes):
        points.append(PointStruct(
            id=str(uuid.uuid4()),
            vector={
                "dense": all_embeddings[i],
                # "sparse" — зарезервировано для будущего hybrid (bge-m3 sparse)
                # В v3 мы используем BM25 на лету, без хранения sparse в Qdrant.
            },
            payload=node,
        ))

    # Batch upsert
    upsert_batch = 128
    for i in range(0, len(points), upsert_batch):
        client.upsert(
            collection_name=settings.QDRANT_COLLECTION,
            points=points[i:i + upsert_batch],
        )

    snapshot_name = create_snapshot()

    parent_count = sum(1 for n in nodes if n["is_parent"])
    child_count = sum(1 for n in nodes if not n["is_parent"])

    logger.info("Ingested '%s': %d parents + %d children = %d total nodes (skipped %d duplicates).",
                doc_name, parent_count, child_count, len(nodes), skipped_duplicates)

    # Обновляем BM25-индекс (для hybrid search в v3)
    try:
        from app.bm25_index import get_bm25_index
        bm25 = get_bm25_index()
        bm25.rebuild_from_qdrant(client, settings.QDRANT_COLLECTION)
        logger.info("BM25 index rebuilt after ingest.")
    except Exception as exc:
        logger.warning("BM25 index rebuild failed (non-blocking): %s", exc)

    return {
        "doc_id": doc_id,
        "doc_name": doc_name,
        "pages": len(pages),
        "parents": parent_count,
        "chunks": child_count,
        "total_nodes": len(nodes),
        "snapshot": snapshot_name,
        "skipped_duplicates": skipped_duplicates,
    }


def _get_existing_chunk_hashes(client: QdrantClient, doc_id: str) -> set:
    """Возвращает множество chunk_hash уже сохранённых чанков документа."""
    existing: set = set()
    try:
        offset = None
        while True:
            result = client.scroll(
                collection_name=settings.QDRANT_COLLECTION,
                limit=500,
                offset=offset,
                with_payload=True,
                with_vectors=False,
                scroll_filter=Filter(
                    must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
                ),
            )
            points, next_offset = result
            if not points:
                break
            for p in points:
                h = (p.payload or {}).get("chunk_hash")
                if h:
                    existing.add(h)
            offset = next_offset
            if not next_offset:
                break
    except Exception as exc:
        logger.warning("Could not fetch existing chunk_hash set: %s", exc)
    return existing


def delete_document(doc_id: str, doc_version: Optional[str] = None) -> bool:
    """Удаляет документ из Qdrant по doc_id (опционально по версии)."""
    client = get_qdrant_client()
    conditions = [FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
    if doc_version:
        conditions.append(FieldCondition(key="doc_version", match=MatchValue(value=doc_version)))
    client.delete(
        collection_name=settings.QDRANT_COLLECTION,
        points_selector=Filter(must=conditions),
    )
    logger.info("Deleted document %s (version=%s)", doc_id, doc_version or "any")

    # Пересобираем BM25-индекс после удаления
    try:
        from app.bm25_index import get_bm25_index
        bm25 = get_bm25_index()
        bm25.rebuild_from_qdrant(client, settings.QDRANT_COLLECTION)
        logger.info("BM25 index rebuilt after delete.")
    except Exception as exc:
        logger.warning("BM25 index rebuild failed (non-blocking): %s", exc)

    return True


def list_documents() -> List[Dict[str, Any]]:
    """Сводка по всем проиндексированным документам (только parent-точки)."""
    client = get_qdrant_client()
    seen: Dict[str, Dict[str, Any]] = {}
    offset: Optional[str] = None

    while True:
        result = client.scroll(
            collection_name=settings.QDRANT_COLLECTION,
            limit=500,
            offset=offset,
            with_payload=True,
            with_vectors=False,
        )
        points, next_offset = result
        if not points:
            break
        for point in points:
            p = point.payload or {}
            doc_id = p.get("doc_id")
            if not doc_id:
                continue
            if doc_id not in seen:
                seen[doc_id] = {
                    "doc_id": doc_id,
                    "doc_name": p.get("doc_name", ""),
                    "doc_version": p.get("doc_version", "v1.0"),
                    "chunks": 0,
                    "pages_set": set(),
                    "section_titles": set(),
                }
            seen[doc_id]["chunks"] += 1
            page = p.get("page_number")
            if page is not None:
                seen[doc_id]["pages_set"].add(page)
            section = p.get("section_title")
            if section:
                seen[doc_id]["section_titles"].add(section)
        offset = next_offset
        if not next_offset:
            break

    return [
        {
            "doc_id": v["doc_id"],
            "doc_name": v["doc_name"],
            "doc_version": v["doc_version"],
            "chunks": v["chunks"],
            "pages": len(v["pages_set"]),
            "sections": sorted(s for s in v["section_titles"] if s)[:5],
        }
        for v in seen.values()
    ]


def fetch_parent_text(parent_id: str) -> Optional[str]:
    """
    Для PDR: возвращает текст parent-чанка по его chunk_id.
    Использует scroll с фильтром по chunk_id.
    """
    client = get_qdrant_client()
    result = client.scroll(
        collection_name=settings.QDRANT_COLLECTION,
        limit=1,
        scroll_filter=Filter(
            must=[FieldCondition(key="chunk_id", match=MatchValue(value=parent_id))]
        ),
        with_payload=True,
        with_vectors=False,
    )
    points, _ = result
    if points:
        return (points[0].payload or {}).get("text")
    return None
